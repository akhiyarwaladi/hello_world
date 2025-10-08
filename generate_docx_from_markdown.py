"""
Auto-generate .docx files from comprehensive markdown documents
Converts Laporan Kemajuan and JICEST Paper markdown to formatted .docx

Author: Claude Code
Date: 2025-10-08
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

def parse_markdown_table(table_text):
    """Parse markdown table into list of lists"""
    lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]

    # Remove separator line (contains |---|---|)
    lines = [line for line in lines if not re.match(r'\|[\s\-:]+\|', line)]

    table_data = []
    for line in lines:
        # Split by | and remove empty first/last elements
        cells = [cell.strip() for cell in line.split('|')]
        cells = [cell for cell in cells if cell]  # Remove empty cells
        if cells:
            table_data.append(cells)

    return table_data

def add_markdown_table_to_doc(doc, table_text, caption=None):
    """Add a formatted table from markdown to the document"""
    table_data = parse_markdown_table(table_text)

    if not table_data:
        return

    # Add caption if provided
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create table
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Light Grid Accent 1'

    # Fill table
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]

            # Clean cell text (remove ** for bold markers)
            cell_text = cell_data.replace('**', '').replace('*', '')
            cell.text = cell_text

            # Header row formatting
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # Data rows
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

                # Center-align numeric columns
                if j > 0 and any(char.isdigit() or char == '.' or char == '%' for char in cell_text):
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add spacing after table
    doc.add_paragraph()

def convert_markdown_to_docx(markdown_file, output_file, title):
    """Convert markdown file to formatted .docx document"""

    # Read markdown content
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Split content into sections
    lines = content.split('\n')

    i = 0
    current_table_lines = []
    in_table = False
    in_code_block = False

    while i < len(lines):
        line = lines[i]

        # Skip YAML front matter or metadata
        if i < 10 and line.startswith('---'):
            i += 1
            continue

        # Code block detection
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            i += 1
            continue

        # Table detection
        if '|' in line and not in_table:
            in_table = True
            current_table_lines = []

        if in_table:
            if '|' in line:
                current_table_lines.append(line)
                i += 1
                continue
            else:
                # End of table
                in_table = False
                if current_table_lines:
                    table_text = '\n'.join(current_table_lines)
                    add_markdown_table_to_doc(doc, table_text)
                    current_table_lines = []

        # Skip empty lines after processing table
        if not line.strip():
            i += 1
            continue

        # Heading detection
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)

        # Bold section titles (e.g., **Tabel 1.**)
        elif line.startswith('**') and '**' in line[2:]:
            p = doc.add_paragraph()
            run = p.add_run(line.replace('**', ''))
            run.bold = True

        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')

        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            doc.add_paragraph(text, style='List Number')

        # Regular paragraph
        elif line.strip():
            # Skip lines that are just markdown formatting
            if not line.startswith('**Tabel') and not line.startswith('|'):
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        i += 1

    # Save document
    doc.save(output_file)
    print(f"[OK] Generated: {output_file}")
    print(f"   File size: {Path(output_file).stat().st_size / 1024:.1f} KB")

def main():
    """Main function to generate both documents"""

    print("="*60)
    print("DOCX DOCUMENT GENERATION FROM MARKDOWN")
    print("="*60)
    print()

    # Define paths
    luaran_dir = Path("luaran")

    # Documents to generate
    documents = [
        {
            'markdown': luaran_dir / "Laporan_Kemajuan_FINAL_WITH_TABLES.md",
            'output': luaran_dir / "Laporan_Kemajuan_Malaria_Detection_UPDATED.docx",
            'title': "LAPORAN KEMAJUAN PENELITIAN\nSISTEM DETEKSI DAN KLASIFIKASI MALARIA"
        },
        # Future: Add JICEST paper when ready
        # {
        #     'markdown': luaran_dir / "JICEST_Paper_FINAL_WITH_TABLES.md",
        #     'output': luaran_dir / "JICEST_Paper_UPDATED.docx",
        #     'title': "Hybrid YOLO-CNN Architecture for Malaria Detection and Classification"
        # }
    ]

    # Generate documents
    for doc_info in documents:
        markdown_file = doc_info['markdown']
        output_file = doc_info['output']
        title = doc_info['title']

        if not markdown_file.exists():
            print(f"[WARNING] Markdown file not found: {markdown_file}")
            continue

        print(f"[Processing] {markdown_file.name}")
        print(f"   Output: {output_file.name}")
        print()

        try:
            convert_markdown_to_docx(markdown_file, output_file, title)
            print()
        except Exception as e:
            print(f"[ERROR] {e}")
            import traceback
            traceback.print_exc()
            print()

    print("="*60)
    print("GENERATION COMPLETE!")
    print("="*60)
    print()
    print("[Generated files]")
    for doc_info in documents:
        output_file = doc_info['output']
        if output_file.exists():
            print(f"   [OK] {output_file}")
            print(f"      Size: {output_file.stat().st_size / 1024:.1f} KB")
    print()
    print("Next steps:")
    print("1. Open the generated .docx files in Microsoft Word")
    print("2. Verify all tables and formatting")
    print("3. Adjust spacing, fonts, or styles as needed")
    print("4. Replace existing .docx files if satisfied")
    print()

if __name__ == "__main__":
    main()
